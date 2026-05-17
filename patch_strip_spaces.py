"""
patch_strip_spaces.py - 移除正文段落中的 ASCII 空格 (U+0020)

只处理普通正文段落，不破坏标题、题注、目录等结构化内容。

跳过规则：
- 标题样式（Heading 1/2/3）
- 题注样式（Caption）
- 目录样式（TOC）
- 英文摘要区域（Abstract → 下一章标题）
- 参考文献区域（参考文献 → 末尾）

用法：python patch_strip_spaces.py <输入文件> [输出文件]
"""

import sys
import os
from docx import Document
from docx.oxml.ns import qn

W = qn

_HEADING_IDS = frozenset(('1', '2', '3', 'Heading1', 'Heading2', 'Heading3',
                          'heading1', 'heading2', 'heading3'))


def _style_id(p):
    """取段落样式 ID，无样式时返回 None"""
    pPr = p.find(W('w:pPr'))
    if pPr is None:
        return None
    ps = pPr.find(W('w:pStyle'))
    return None if ps is None else ps.get(W('w:val'))


def _is_non_body(sid):
    """非正文样式（标题、题注、目录）"""
    if sid is None:
        return False
    if sid in _HEADING_IDS:
        return True
    if 'caption' in sid.lower():
        return True
    if sid.lower().startswith('toc'):
        return True
    return False


def _find_skip_indices(doc):
    """收集需要跳过的段落序号"""
    skip = set()
    in_en_abs = False
    in_bib = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        sid = _style_id(para._element)

        # 样式跳转
        if _is_non_body(sid):
            skip.add(i)

        # ── 英文摘要区域 ──
        if text == 'Abstract' and not in_en_abs:
            in_en_abs = True
            continue
        if in_en_abs:
            if sid and sid in _HEADING_IDS:
                in_en_abs = False
            else:
                skip.add(i)
                continue

        # ── 参考文献区域 ──
        if text == '参考文献':
            in_bib = True
        if in_bib:
            skip.add(i)
            continue

    return skip


def patch_strip_spaces(input_path: str, output_path: str) -> int:
    doc = Document(input_path)
    skip_indices = _find_skip_indices(doc)

    total = 0
    for i, para in enumerate(doc.paragraphs):
        if i in skip_indices:
            continue

        for r in para._element.iter(W('w:r')):
            for t in r.iter(W('w:t')):
                if not t.text:
                    continue
                parent = t.getparent()
                if parent is not None and parent.tag == W('w:instrText'):
                    continue
                stripped = t.text.replace(' ', '')
                if len(stripped) != len(t.text):
                    t.text = stripped
                    total += 1

    doc.save(output_path)
    body_count = len(doc.paragraphs) - len(skip_indices)
    print(f'  跳过 {len(skip_indices)} 个非正文段落，处理 {body_count} 个正文段落')
    print(f'  已移除 {total} 个空格')
    return total


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else inp
    if not os.path.exists(inp):
        print(f'错误: 文件不存在 {inp}')
        sys.exit(1)
    patch_strip_spaces(inp, out)
