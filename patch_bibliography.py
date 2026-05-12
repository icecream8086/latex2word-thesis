"""
patch_bibliography.py - 修补"参考文献"标题和内容格式

功能（通过控制变量开关）：
  FORMAT_BIBLIOGRAPHY - 标题/字体/缩进格式化
  URL_HYPERLINK       - 将参考文献中的 URL 转为可点击超链接
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import sys
import os
import re

# ==================== 控制变量 ====================
FORMAT_BIBLIOGRAPHY = True
URL_HYPERLINK = True
# =================================================

URL_PATTERN = re.compile(r'(https?://[^\s。，、；：\)）\]]+)')


def set_run_font(run, font_name='SimSun', font_size=Pt(12)):
    """设置 run 的字体和字号"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
        font_name
    )
    run.font.size = font_size


def set_indent_zero(paragraph):
    """将段落的首行缩进设为0"""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = paragraph._element.makeelement(qn('w:pPr'), {})
        paragraph._element.insert(0, pPr)

    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = pPr.makeelement(qn('w:ind'), {})
        pPr.append(ind)

    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')
    return True


def _make_hyperlink_run(p_elem, url, font_name='SimSun', font_size=Pt(12)):
    """创建一个超链接 run 元素（带蓝色+下划线样式）"""
    hyperlink = p_elem.makeelement(qn('w:hyperlink'), {
        qn('w:history'): '1'
    })

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 字体
    rFonts = rPr.makeelement(qn('w:rFonts'), {})
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

    # 字号
    sz = rPr.makeelement(qn('w:sz'), {})
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    szCs = rPr.makeelement(qn('w:szCs'), {})
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(szCs)

    # 超链接样式：蓝色 + 下划线
    color = rPr.makeelement(qn('w:color'), {})
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)
    u = rPr.makeelement(qn('w:u'), {})
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    r.append(rPr)
    t = r.makeelement(qn('w:t'), {})
    t.text = url
    r.append(t)
    hyperlink.append(r)

    return hyperlink


def make_urls_clickable(doc):
    """在参考文献段落中检测 URL 并转换为可点击超链接"""
    found_refs = False
    modified = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        # 标记参考文献段落（"参考文献"标题之后的段落）
        if text == '参考文献' and not found_refs:
            found_refs = True
            continue

        if not found_refs or not text:
            continue

        # 跳过标题段落
        if text in ('参考文献',):
            continue

        # 查找 URL
        urls = URL_PATTERN.findall(text)
        if not urls:
            continue

        p = paragraph._element
        parts = URL_PATTERN.split(paragraph.text)
        # parts 格式: [before, url1, between, url2, after, ...]

        # 清除段落原有内容
        for child in list(p):
            if child.tag != qn('w:pPr'):
                p.remove(child)

        # 重建段落：普通文本用 run，URL 用 hyperlink
        for i, part in enumerate(parts):
            if not part:
                continue
            if i % 2 == 1:
                # 奇数索引 → URL，创建超链接
                p.append(_make_hyperlink_run(p, part))
            else:
                # 偶数索引 → 普通文本
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')

                rFonts = rPr.makeelement(qn('w:rFonts'), {})
                rFonts.set(qn('w:ascii'), 'SimSun')
                rFonts.set(qn('w:hAnsi'), 'SimSun')
                rFonts.set(qn('w:eastAsia'), 'SimSun')
                rPr.append(rFonts)

                sz = rPr.makeelement(qn('w:sz'), {})
                sz.set(qn('w:val'), '24')
                rPr.append(sz)
                szCs = rPr.makeelement(qn('w:szCs'), {})
                szCs.set(qn('w:val'), '24')
                rPr.append(szCs)

                r.append(rPr)
                t = r.makeelement(qn('w:t'), {})
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t.text = part
                r.append(t)
                p.append(r)

        modified += 1

    if modified:
        print(f'  [超链接] 已将 {modified} 个参考文献中的 URL 转为可点击超链接')
    return modified


def _format_as_chapter_title(paragraph, doc, display_text):
    """将段落格式化为章级标题：Heading 1 + 黑体小二居中分页"""
    paragraph.clear()
    run = paragraph.add_run(display_text)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'SimHei'
    )
    run.font.size = Pt(18)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_indent_zero(paragraph)

    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = paragraph._element.makeelement(qn('w:pPr'), {})
        paragraph._element.insert(0, pPr)
    pb = pPr.find(qn('w:pageBreakBefore'))
    if pb is None:
        pPr.append(pPr.makeelement(qn('w:pageBreakBefore'), {}))

    paragraph.style = doc.styles['heading 1']
    # 清理编号和 outline 覆盖
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
        outLvl = pPr.find(qn('w:outlineLvl'))
        if outLvl is not None:
            pPr.remove(outLvl)

    print(f'  已修改: "{display_text}" (黑体, 小二, 居中, 分页, 目录章级)')


def patch_bibliography(input_path, output_path):
    doc = Document(input_path)
    found_refs = False
    abstract_done = False
    title_modified = False
    content_modified = False
    ref_counter = 1

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        # 查找 "Abstract" 标题（在文档开头，早于参考文献）
        if text.replace(' ', '') == 'Abstract' and not abstract_done and FORMAT_BIBLIOGRAPHY:
            print(f'找到"Abstract"段落: "{paragraph.text}"')
            _format_as_chapter_title(paragraph, doc, text)
            abstract_done = True
            title_modified = True
            continue

        # 查找"参考文献"标题
        if text == '参考文献' and not found_refs:
            found_refs = True

            if not FORMAT_BIBLIOGRAPHY:
                continue

            _format_as_chapter_title(paragraph, doc, '参考文献')
            title_modified = True
            continue

        # 处理参考文献内容段落
        if found_refs and text.strip():
            if not FORMAT_BIBLIOGRAPHY:
                continue

            if not text.startswith('['):
                new_text = f'[{ref_counter}] {text}'
                ref_counter += 1

                paragraph.clear()
                run = paragraph.add_run(new_text)
                set_run_font(run, 'SimSun', Pt(12))
            else:
                for run in paragraph.runs:
                    set_run_font(run, 'SimSun', Pt(12))
                ref_counter += 1

            set_indent_zero(paragraph)
            content_modified = True
            current_text = paragraph.text.strip()
            print(f'  设置字体/去缩进: "{current_text[:40]}..."')

    # URL → 超链接（独立阶段，在格式化之后执行）
    if URL_HYPERLINK:
        make_urls_clickable(doc)

    if title_modified or content_modified or URL_HYPERLINK:
        doc.save(output_path)
        print(f'处理完成，已保存至: {output_path}')
    else:
        print('未找到"参考文献"段落，无需修改')
        doc.save(output_path)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python patch_bibliography.py <输入文件> [输出文件]')
        print()
        print('控制变量（在脚本顶部修改）:')
        print(f'  FORMAT_BIBLIOGRAPHY = {FORMAT_BIBLIOGRAPHY}')
        print(f'  URL_HYPERLINK       = {URL_HYPERLINK}')
        sys.exit(1)

    input_file = sys.argv[1]
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        base, ext = os.path.splitext(input_file)
        output_file = f'{base}_bibl{ext}'

    if not os.path.exists(input_file):
        print(f'错误: 输入文件不存在 - {input_file}')
        sys.exit(1)

    patch_bibliography(input_file, output_file)
