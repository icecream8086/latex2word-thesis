"""
patch_citation_hyperlink.py - 参考文献引用标记与正文之间的双向超链接

功能：
  1. 将正文中的引用标记 [N] 转换为指向参考文献列表条目的超链接（正向）
  2. 在首次出现 [N] 的正文位置添加书签 _BackToCite_N
  3. 将参考文献条目中的 [N] 转换为指向正文首次引用位置的超链接（反向）
  4. 为参考文献条目添加书签 _CiteRef_N

依赖：python-docx
用法：python patch_citation_hyperlink.py <输入文件> [输出文件]
"""

import sys
import os
import re
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W = qn
CITE_PAT = re.compile(r'\[(\d+(?:\s*[, \-]\s*\d+)*)\]')
_BM_COUNTER = [0]  # 全局书签 ID 计数器


def _next_bm_id() -> str:
    _BM_COUNTER[0] += 1
    return str(_BM_COUNTER[0])


def _add_bookmark(p, name):
    """为段落添加书签（保留现有 ID）"""
    ids = set()
    for bs in p.iter(W('w:bookmarkStart')):
        try:
            ids.add(int(bs.get(W('w:id'), '0')))
        except ValueError:
            pass
    bid = str(max(ids) + 1 if ids else 0)

    bms = OxmlElement('w:bookmarkStart')
    bms.set(W('w:id'), bid)
    bms.set(W('w:name'), name)
    bme = OxmlElement('w:bookmarkEnd')
    bme.set(W('w:id'), bid)

    fr = p.find(W('w:r'))
    if fr is not None:
        fr.addprevious(bms)
    else:
        pPr = p.find(W('w:pPr'))
        if pPr is not None:
            pPr.addnext(bms)
        else:
            p.insert(0, bms)
    p.append(bme)


def _make_hyperlink(anchor, text, rpr=None):
    """创建 w:hyperlink > w:r > w:t（自动清除蓝色+下划线格式）"""
    hl = OxmlElement('w:hyperlink')
    hl.set(W('w:anchor'), anchor)
    hl.set(W('w:history'), '1')
    r = OxmlElement('w:r')

    if rpr is not None:
        clean_rpr = copy.deepcopy(rpr)
        for tag in ('w:color', 'w:u', 'w:uColor'):
            for el in clean_rpr.findall(W(tag)):
                clean_rpr.remove(el)
        color = OxmlElement('w:color')
        color.set(W('w:val'), 'auto')
        clean_rpr.append(color)
        r.append(clean_rpr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    hl.append(r)
    return hl


def _make_hyperlink_with_bookmark(anchor, back_name, text, rpr=None):
    """创建 hyperlink 并包裹书签：用于首次引用（正文→文献、文献→正文双向锚点）"""
    hl = _make_hyperlink(anchor, text, rpr)
    if back_name is None:
        return hl

    bid = _next_bm_id()
    bms = OxmlElement('w:bookmarkStart')
    bms.set(W('w:id'), bid)
    bms.set(W('w:name'), back_name)
    bme = OxmlElement('w:bookmarkEnd')
    bme.set(W('w:id'), bid)

    # 返回容器：书签 start + hyperlink + 书签 end
    container = OxmlElement('w:r')  # 仅作为占位容器
    container.append(bms)
    container.append(hl)
    container.append(bme)
    return container


def _make_run(text, rpr=None):
    """创建 w:r > w:t"""
    r = OxmlElement('w:r')
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


# ── 正文引用 → 参考文献（正向） ──

def add_entry_bookmarks(doc):
    """为参考文献列表中的每条 [N] 条目添加书签 _CiteRef_N"""
    in_bib = False
    n = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == '参考文献':
            in_bib = True
            continue
        if not in_bib or not text:
            continue
        m = re.match(r'^\[(\d+)\]', text)
        if m:
            _add_bookmark(para._element, f'_CiteRef_{m.group(1)}')
            n += 1
    return n


def _convert_para(p, first_seen):
    """
    将单个段落中的引用标记转为超链接。
    首次出现的引用额外包裹 _BackToCite_N 书签。
    返回 True 表示有修改。
    """
    children = list(p.iterchildren())

    runs_info = []
    pos = 0
    for ci, child in enumerate(children):
        if child.tag == W('w:r'):
            t = child.find(W('w:t'))
            if t is not None and t.text:
                runs_info.append((ci, child, t.text, pos))
                pos += len(t.text)

    if not runs_info:
        return False

    full = ''.join(t for _, _, t, _ in runs_info)
    matches = list(CITE_PAT.finditer(full))
    if not matches:
        return False

    cite_of = {}
    for m in matches:
        first_num = int(re.search(r'\d+', m.group(1)).group())
        for ppos in range(m.start(), m.end()):
            cite_of[ppos] = first_num

    new_children = []
    for ci, child in enumerate(children):
        rinfo = next((ri for ri in runs_info if ri[0] == ci), None)
        if rinfo is None:
            new_children.append(child)
            continue

        _, elem, text, gs = rinfo
        rpr = elem.find(W('w:rPr'))

        has_cite = any((gs + local_ci) in cite_of for local_ci in range(len(text)))
        if not has_cite:
            new_children.append(elem)
            continue

        segs = []
        cur_type = None
        cur_chars = []
        cur_num = None

        for local_ci, ch in enumerate(text):
            gpos = gs + local_ci
            in_cite = gpos in cite_of
            this_type = ('cite', cite_of[gpos]) if in_cite else ('text', None)

            if this_type != cur_type and cur_chars:
                segs.append((cur_type, ''.join(cur_chars), cur_num))
                cur_chars = []

            cur_type = this_type
            cur_chars.append(ch)
            if in_cite:
                cur_num = cite_of[gpos]

        if cur_chars:
            segs.append((cur_type, ''.join(cur_chars), cur_num))

        for seg_type, seg_text, num in segs:
            if not seg_text:
                continue
            if seg_type[0] == 'text':
                new_children.append(_make_run(seg_text, rpr))
            else:
                anchor = f'_CiteRef_{num}'
                back_name = None
                if num not in first_seen:
                    first_seen.add(num)
                    back_name = f'_BackToCite_{num}'

                container = _make_hyperlink_with_bookmark(anchor, back_name, seg_text, rpr)
                # 展开容器中的子元素
                for sub in list(container):
                    container.remove(sub)
                    new_children.append(sub)

    for child in list(p):
        p.remove(child)
    for child in new_children:
        p.append(child)

    return True


def convert_para_citations(doc, first_seen):
    """将正文段落中引用标记转为超链接，首次引用加书签"""
    in_bib = False
    n = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == '参考文献':
            in_bib = True
        if in_bib or not text:
            continue
        if not CITE_PAT.search(para.text):
            continue
        if _convert_para(para._element, first_seen):
            n += 1
    return n


# ── 参考文献 → 正文（反向） ──

def add_bibliography_backlinks(doc, first_seen):
    """
    将参考文献条目开头的 [N] 转换为指向正文首次引用位置的超链接。
    需要 first_seen 集合得知哪些编号在正文中被引用过。
    """
    in_bib = False
    count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text == '参考文献':
            in_bib = True
            continue
        if not in_bib or not text:
            continue

        m = re.match(r'^(\[(\d+)\])', text)
        if not m:
            continue

        num = int(m.group(2))
        back_name = f'_BackToCite_{num}'
        cite_tag = m.group(1)

        # 该编号在正文中从未被引用 → 无法建立反向链接
        if num not in first_seen:
            continue

        p = para._element
        for r in p.iter(W('w:r')):
            t = r.find(W('w:t'))
            if t is None or not t.text:
                continue
            if cite_tag not in t.text:
                continue

            # 找到包含 [N] 的 run → 分割
            run_text = t.text
            idx = run_text.index(cite_tag)
            before = run_text[:idx]
            after = run_text[idx + len(cite_tag):]
            rpr = r.find(W('w:rPr'))

            parent = r.getparent()
            if parent is None:
                continue
            pos_in_parent = list(parent).index(r)

            replacements = []
            if before:
                replacements.append(_make_run(before, rpr))
            replacements.append(_make_hyperlink(back_name, cite_tag, rpr))
            if after:
                replacements.append(_make_run(after, rpr))

            parent.remove(r)
            for j, repl in enumerate(replacements):
                parent.insert(pos_in_parent + j, repl)

            count += 1
            break  # 一个条目只处理一次

    return count


# ── 主入口 ──

def patch_citation_hyperlink(input_path, output_path):
    """主函数"""
    doc = Document(input_path)

    # 第一步：正文引用 → 超链接 + 书签
    first_seen = set()
    cn = convert_para_citations(doc, first_seen)
    print(f'  已转换 {cn} 个段落中的引用标记为超链接')

    # 第二步：参考文献条目书签
    bm = add_entry_bookmarks(doc)
    print(f'  已为 {bm} 个参考文献条目添加书签')

    # 第三步：参考文献 → 正文反向链接
    bl = add_bibliography_backlinks(doc, first_seen)
    print(f'  已添加 {bl} 个参考文献反向超链接')

    doc.save(output_path)
    print(f'  已保存: {output_path}')
    return cn


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else inp
    if not os.path.exists(inp):
        print(f'错误: 文件不存在 {inp}')
        sys.exit(1)
    patch_citation_hyperlink(inp, out)
