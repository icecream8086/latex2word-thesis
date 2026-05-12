"""
patch_thanks.py - 修补"致谢""摘要""目录"标题格式

将"致谢""摘要""目录"标题改为：
- 字体：黑体（SimHei）
- 字号：小二（18pt）
- 加粗、居中
- 去掉首行缩进、分页
- Heading 1 样式（出现在目录中，但不带章节编号）
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import sys
import os
import re

# ==================== 控制变量 ====================
TITLE_SPACING = '    '    # 标题字间距（汉字之间的空格）
FORMAT_TOC = True         # 格式化"目录"标题
# =================================================


def _strip_spaces(text):
    """去掉文本中所有空白（ASCII + 全角空格）"""
    return re.sub(r'[\s　]', '', text)


def _format_as_chapter_title(paragraph, doc, display_text):
    """将段落格式化为章级标题：Heading 1 + 黑体小二居中分页"""
    paragraph.clear()
    run = paragraph.add_run(display_text)
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'SimHei'
    )
    run.font.size = Pt(18)
    run.bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)

    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = paragraph._element.makeelement(qn('w:pPr'), {})
        paragraph._element.insert(0, pPr)

    # 分页
    pb = pPr.find(qn('w:pageBreakBefore'))
    if pb is None:
        pPr.append(pPr.makeelement(qn('w:pageBreakBefore'), {}))

    # 删除段落级 outlineLvl 覆盖（恢复为样式继承值）
    ol = pPr.find(qn('w:outlineLvl'))
    if ol is not None:
        pPr.remove(ol)
    # 删除编号
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)

    # 设为 Heading 1 样式
    paragraph.style = doc.styles['heading 1']

    print(f'  已修改为: "{display_text}" (黑体, 小二, 居中, 分页, 目录章级)')


def patch_thanks(input_path, output_path):
    doc = Document(input_path)
    modified = False
    thanks_done = False
    abstract_done = False
    toc_done = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        cleaned = _strip_spaces(text)

        # 匹配"致谢"（中间可能有空格）
        if cleaned == '致谢' and not thanks_done:
            print(f'找到"致谢"段落: "{paragraph.text}"')
            _format_as_chapter_title(paragraph, doc, f'致{TITLE_SPACING}谢')
            thanks_done = True
            modified = True
            continue

        # 匹配"摘要"（中间可能有空格）
        if cleaned == '摘要' and not abstract_done:
            print(f'找到"摘要"段落: "{paragraph.text}"')
            _format_as_chapter_title(paragraph, doc, f'摘{TITLE_SPACING}要')
            abstract_done = True
            modified = True
            continue

        # 匹配"目录"（中间可能有空格）
        if FORMAT_TOC and cleaned == '目录' and not toc_done:
            print(f'找到"目录"段落: "{paragraph.text}"')
            _format_as_chapter_title(paragraph, doc, f'目{TITLE_SPACING}录')
            toc_done = True
            modified = True
            continue

    if modified:
        doc.save(output_path)
        print(f'处理完成，已保存至: {output_path}')
    else:
        print('未找到需要处理的段落，无需修改')
        doc.save(output_path)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python patch_thanks.py <输入文件> [输出文件]')
        sys.exit(1)

    input_file = sys.argv[1]
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        base, ext = os.path.splitext(input_file)
        output_file = f'{base}_thanks{ext}'

    if not os.path.exists(input_file):
        print(f'错误: 输入文件不存在 - {input_file}')
        sys.exit(1)

    patch_thanks(input_file, output_file)
