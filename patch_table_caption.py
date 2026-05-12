"""
patch_table_caption.py - 文档后处理工具（纯 python-docx，无 COM 依赖）

支持的操作（通过文件顶部的控制变量开关）：
  BOLD_CAPTION       - 表格标题 (Table Caption) 加粗
  RESOLVE_HYPERLINKS - 将 HYPERLINK 域代码 (fldChar+instrText) 转换为 w:hyperlink
  AUTO_FIT_TABLES    - 表格宽度自适应页面
"""

from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import sys
import os
import re

# ==================== 控制变量 ====================
BOLD_CAPTION = True
RESOLVE_HYPERLINKS = True
AUTO_FIT_TABLES = False
# =================================================


def _bold_run_element(r_elem):
    """给单个 <w:r> 元素设置加粗 (添加 <w:b/> 到 rPr)"""
    rPr = r_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r_elem.insert(0, rPr)
    if rPr.find(qn("w:b")) is None:
        bold = OxmlElement("w:b")
        rPr.append(bold)


def fix_caption_bold(doc):
    """
    将 Table Caption 段落中所有文本（含超链接内文本）设为加粗。
    非破坏性操作 —— 保留原有 run 结构和超链接。
    """
    modified = 0
    for para in doc.paragraphs:
        if para.style.name != "Table Caption":
            continue
        p = para._element
        for r_elem in p.iter(qn("w:r")):
            _bold_run_element(r_elem)
        modified += 1
    if modified:
        print(f"  [加粗] 已处理 {modified} 个表格标题")
    return modified


def resolve_hyperlink_fields(doc):
    """
    将文档中的 HYPERLINK 域代码 (fldChar + instrText) 转换为 w:hyperlink 元素。

    转换前:
      <w:p>
        <w:r><w:fldChar w:fldCharType="begin"/></w:r>
        <w:r><w:instrText>HYPERLINK \\l "bookmark"</w:instrText></w:r>
        <w:r><w:fldChar w:fldCharType="separate"/></w:r>
        <w:r ...>result text</w:r>
        <w:r><w:fldChar w:fldCharType="end"/></w:r>
        <w:r ...>more text after field</w:r>
      </w:p>

    转换后:
      <w:p>
        <w:hyperlink w:anchor="bookmark">
          <w:r ...>result text</w:r>
        </w:hyperlink>
        <w:r ...>more text after field</w:r>
      </w:p>
    """
    resolved = 0

    for para in doc.paragraphs:
        p = para._element
        children = list(p)

        begin = sep = end = None
        bookmark = None

        for i, child in enumerate(children):
            if child.tag != qn("w:r"):
                continue

            fc = child.find(qn("w:fldChar"))
            if fc is not None:
                t = fc.get(qn("w:fldCharType"))
                if t == "begin":
                    begin = i
                elif t == "separate":
                    sep = i
                elif t == "end":
                    end = i
                    break
                continue

            if begin is not None and sep is None:
                it = child.find(qn("w:instrText"))
                if it is not None and (it.text or "").strip().startswith("HYPERLINK"):
                    m = re.search(r'\\(?:l)\s+"([^"]+)"', it.text or "")
                    if m:
                        bookmark = m.group(1)

        if begin is None or end is None or sep is None or bookmark is None:
            continue

        # Runs between 'separate' and 'end' form the link text
        result_runs = children[sep + 1 : end]

        hyperlink = p.makeelement(qn("w:hyperlink"), {qn("w:anchor"): bookmark})
        for run in result_runs:
            hyperlink.append(run)

        # Insert hyperlink at the position of the 'begin' marker
        begin_elem = children[begin]
        p.insert(list(p).index(begin_elem), hyperlink)

        # Remove old field code elements (begin, instr, separate, end)
        for elem in children[begin : end + 1]:
            if elem.getparent() is not None:
                try:
                    p.remove(elem)
                except ValueError:
                    pass

        resolved += 1

    if resolved:
        print(f"  [超链接] 已解析 {resolved} 个 HYPERLINK 域代码")
    return resolved


def auto_fit_tables(doc):
    """设置所有表格宽度为自适应页面"""
    modified = 0
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
            tbl.insert(0, tblPr)

        old_w = tblPr.find(qn("w:tblW"))
        if old_w is not None:
            tblPr.remove(old_w)
        tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>'))

        old_layout = tblPr.find(qn("w:tblLayout"))
        if old_layout is not None:
            tblPr.remove(old_layout)
        tblPr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="autofit"/>'))

        modified += 1

    if modified:
        print(f"  [表格] 已调整 {modified} 个表格宽度")
    return modified


def process_document(input_path, output_path):
    """按控制变量设置，按合理顺序执行所有启用的操作"""
    print(f"处理文档: {input_path}")
    doc = Document(input_path)

    # 1. 先解析域代码（否则后面合并 run 会丢失域代码结构）
    if RESOLVE_HYPERLINKS:
        resolve_hyperlink_fields(doc)

    # 2. 再处理加粗（非破坏性，保留超链接等结构）
    if BOLD_CAPTION:
        fix_caption_bold(doc)

    # 3. 最后调整表格宽度
    if AUTO_FIT_TABLES:
        auto_fit_tables(doc)

    doc.save(output_path)
    print(f"完成，已保存至: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python patch_table_caption.py <输入文件> [输出文件]")
        print()
        print("控制变量（在脚本顶部修改）:")
        print(f"  BOLD_CAPTION       = {BOLD_CAPTION}")
        print(f"  RESOLVE_HYPERLINKS = {RESOLVE_HYPERLINKS}")
        print(f"  AUTO_FIT_TABLES    = {AUTO_FIT_TABLES}")
        sys.exit(1)

    input_file = sys.argv[1]
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        base, ext = os.path.splitext(input_file)
        output_file = f"{base}_patched{ext}"

    if not os.path.exists(input_file):
        print(f"错误: 输入文件不存在 - {input_file}")
        sys.exit(1)

    process_document(input_file, output_file)
