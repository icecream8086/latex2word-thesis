"""
patch_pagenum.py - 为 Word 文档添加不同区域的页码（纯 python-docx 实现）

功能：
- 正文前（摘要、目录）：使用大写罗马数字（Ⅰ、Ⅱ、Ⅲ），从1开始
- 正文及之后：使用阿拉伯数字（1、2、3），从正文第1页开始

用法：python patch_pagenum.py <输入文件> [输出文件]

依赖：python-docx
"""

import os
import sys
import re
import zipfile
import tempfile
from copy import deepcopy
from lxml import etree
from docx import Document

# ── 命名空间 ──
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NS_R = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
NS_REL = 'http://schemas.openxmlformats.org/package/2006/relationships'
NS_CT = 'http://schemas.openxmlformats.org/package/2006/content-types'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

FOOTER_CT = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml'
FOOTER_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer'

_H1_KEYWORDS = ('heading 1', '标题 1', 'heading1')
_FRONT_MATTER = ('摘要', 'abstract')


# ── 通用 XML 辅助 ──

def _W(tag):
    return f'{{{NS_W}}}{tag}'

def _R(tag):
    return f'{{{NS_R}}}{tag}'

def _clean_text(raw):
    return raw.replace('\r', '').replace('\x07', '').replace('\x0b', '').strip()

def _is_front_matter(text):
    return re.sub(r'[\s　 ]', '', text).lower() in _FRONT_MATTER


# ── 构建页脚 XML ──

def _build_footer_xml(centered=True, empty=False):
    """返回 bytes 的 <w:ftr> XML"""
    W = _W
    root = etree.Element(W('ftr'))

    if empty:
        p = etree.SubElement(root, W('p'))
        pPr = etree.SubElement(p, W('pPr'))
        ps = etree.SubElement(pPr, W('pStyle'))
        ps.set(W('val'), 'af2')
        return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # 第一段：PAGE 域
    p1 = etree.SubElement(root, W('p'))
    pPr1 = etree.SubElement(p1, W('pPr'))

    etree.SubElement(pPr1, W('pStyle')).set(W('val'), 'af2')

    fp = etree.SubElement(pPr1, W('framePr'))
    fp.set(W('wrap'), 'around')
    fp.set(W('vAnchor'), 'text')
    fp.set(W('hAnchor'), 'margin')
    fp.set(W('xAlign'), 'center')
    fp.set(W('y'), '1')

    if centered:
        etree.SubElement(pPr1, W('jc')).set(W('val'), 'center')

    # <w:rPr><w:rStyle w:val="af5"/></w:rPr> （在 pPr 中）
    rpr = etree.SubElement(pPr1, W('rPr'))
    etree.SubElement(rpr, W('rStyle')).set(W('val'), 'af5')

    # PAGE 域
    def _run():
        r = etree.Element(W('r'))
        rpr2 = etree.SubElement(r, W('rPr'))
        etree.SubElement(rpr2, W('rStyle')).set(W('val'), 'af5')
        return r

    # begin
    r0 = _run()
    etree.SubElement(r0, W('fldChar')).set(W('fldCharType'), 'begin')
    p1.append(r0)

    # instrText
    r1 = _run()
    instr = etree.SubElement(r1, W('instrText'))
    instr.set(XML_SPACE, 'preserve')
    instr.text = ' PAGE '
    p1.append(r1)

    # separate
    r2 = _run()
    etree.SubElement(r2, W('fldChar')).set(W('fldCharType'), 'separate')
    p1.append(r2)

    # end
    r3 = _run()
    etree.SubElement(r3, W('fldChar')).set(W('fldCharType'), 'end')
    p1.append(r3)

    # 第二段：空行
    p2 = etree.SubElement(root, W('p'))
    pPr2 = etree.SubElement(p2, W('pPr'))
    etree.SubElement(pPr2, W('pStyle')).set(W('val'), 'af2')

    return etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


# ── 包级操作 ──

def _read_zip(path):
    """将 docx 读为 dict[name] = bytes"""
    with zipfile.ZipFile(path, 'r') as z:
        return {n: z.read(n) for n in z.namelist()}

def _write_zip(path, content):
    """将 dict[name] = bytes 写为 docx"""
    with zipfile.ZipFile(path, 'w', zipfile.ZIP_DEFLATED) as z:
        for name in sorted(content):
            z.writestr(name, content[name])

def _next_rid(rels_root):
    """找下一个可用 rId"""
    used = set()
    for rel in rels_root:
        rid = rel.get('Id', '')
        if rid.startswith('rId'):
            used.add(rid)
    n = 1
    while f'rId{n}' in used:
        n += 1
    return f'rId{n}'

def _next_footer_num(content):
    """找下一个可用页脚编号"""
    n = 1
    while f'word/footer{n}.xml' in content:
        n += 1
    return n

def _update_ct(content, part_name, content_type):
    """在 [Content_Types].xml 中添加 Override"""
    ct = content.get('[Content_Types].xml', b'')
    if not ct:
        return
    root = etree.fromstring(ct)
    # 检查是否已有
    for ov in root:
        if ov.get('PartName', '') == f'/{part_name}':
            return  # 已存在
    ov = etree.SubElement(root, f'{{{NS_CT}}}Override')
    ov.set('PartName', f'/{part_name}')
    ov.set('ContentType', content_type)
    # 按 PartName 排序
    children = sorted(root, key=lambda x: x.get('PartName', ''))
    for c in children:
        root.remove(c)
    for c in children:
        root.append(c)
    content['[Content_Types].xml'] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)


def _clean_old_footers(content):
    """清除所有旧的页脚文件和引用"""
    footer_files = [n for n in content if n.startswith('word/footer') and n.endswith('.xml')]

    # 从关系文件中移除
    rels_path = 'word/_rels/document.xml.rels'
    if rels_path in content:
        root = etree.fromstring(content[rels_path])
        to_remove = []
        for rel in root:
            if rel.get('Type') == FOOTER_REL_TYPE:
                to_remove.append(rel)
        for rel in to_remove:
            root.remove(rel)
        content[rels_path] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # 从 document.xml 中移除 footerReference
    if 'word/document.xml' in content:
        root = etree.fromstring(content['word/document.xml'])
        for fr in root.iter(_W('footerReference')):
            fr.getparent().remove(fr)
        content['word/document.xml'] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # 从 [Content_Types].xml 移除
    ct_content = content.get('[Content_Types].xml', b'')
    if ct_content:
        root = etree.fromstring(ct_content)
        for ov in list(root):
            pn = ov.get('PartName', '')
            if '/word/footer' in pn:
                root.remove(ov)
        content['[Content_Types].xml'] = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # 删除页脚文件
    for f in footer_files:
        content.pop(f, None)


def _inject_footer(content, section_idx, footer_type, ftr_bytes):
    """
    向指定节注入一个页脚。
    创建 footer XML 文件 → 添加 rel → 添加 footerReference → 更新 CT。
    """
    # 计算新文件编号
    num = _next_footer_num(content)
    fname = f'word/footer{num}.xml'

    # 写入文件
    content[fname] = ftr_bytes

    # 更新 [Content_Types].xml
    _update_ct(content, fname, FOOTER_CT)

    # 更新关系文件
    rels_path = 'word/_rels/document.xml.rels'
    rels_root = etree.fromstring(content[rels_path])
    rid = _next_rid(rels_root)
    rel = etree.SubElement(rels_root, f'{{{NS_REL}}}Relationship')
    rel.set('Id', rid)
    rel.set('Type', FOOTER_REL_TYPE)
    rel.set('Target', f'footer{num}.xml')
    content[rels_path] = etree.tostring(rels_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # 更新 document.xml 中的 footerReference
    doc_root = etree.fromstring(content['word/document.xml'])
    body = doc_root.find(_W('body'))
    all_sect = body.findall(f'.//{_W("sectPr")}')

    if section_idx < len(all_sect):
        sp = all_sect[section_idx]
        # 移除同类型的旧引用
        for fr in sp.findall(_W('footerReference')):
            if fr.get(_W('type')) == footer_type:
                sp.remove(fr)
        # 添加新引用
        fr = etree.SubElement(sp, _W('footerReference'))
        fr.set(_W('type'), footer_type)
        fr.set(_R('id'), rid)

    content['word/document.xml'] = etree.tostring(doc_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    return rid


# ── 节操作 ──

def _find_ch1_start_index(content):
    """
    使用 python-docx 找到第一个非正文前 Heading 1 的段落索引。
    返回 (索引, 文本) 或 (None, None)。
    """
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = tmp.name
    try:
        _write_zip(tmp_path, content)
        doc = Document(tmp_path)
        for i, para in enumerate(doc.paragraphs):
            raw = _clean_text(para.text)
            if not raw:
                continue
            style_name = (para.style.name or '').lower() if para.style else ''
            if not any(kw in style_name for kw in _H1_KEYWORDS):
                continue
            if _is_front_matter(raw):
                continue
            return i, raw
        return None, None
    finally:
        os.unlink(tmp_path)


def _fix_sections_in_package(content, ch1_index=None):
    """
    直接在包 XML 中修复分节符：
    1. 移除所有段落级 sectPr
    2. 保留最后一个 body sectPr
    3. 在 ch1_index 对应段落前插入包含下一页分节符的段落
    """
    doc_root = etree.fromstring(content['word/document.xml'])
    body = doc_root.find(_W('body'))

    # 1. 移除所有段落级 sectPr
    for para in body.findall(_W('p')):
        pPr = para.find(_W('pPr'))
        if pPr is not None:
            for sp in pPr.findall(_W('sectPr')):
                pPr.remove(sp)

    # 2. 保留最后一个 body sectPr
    body_sect_prs = body.findall(_W('sectPr'))
    for sp in body_sect_prs[:-1]:
        body.remove(sp)

    # 获取最终保留的 sectPr（用于复制页设置）
    ref_sect_pr = body.findall(_W('sectPr'))
    ref_sect_pr = ref_sect_pr[-1] if ref_sect_pr else None

    # 3. 在 ch1_index 段落前插入分节符
    if ch1_index is not None:
        all_paras = list(body.iterchildren(_W('p')))

        if ch1_index < len(all_paras):
            ch1_para = all_paras[ch1_index]

            # 移除 pageBreakBefore
            pPr = ch1_para.find(_W('pPr'))
            if pPr is not None:
                pb = pPr.find(_W('pageBreakBefore'))
                if pb is not None:
                    pPr.remove(pb)

            # 创建分节符段落（oddPage 保证每章从奇数页开始）
            new_para_str = f'<w:p xmlns:w="{NS_W}"><w:pPr><w:sectPr w:type="oddPage"/></w:pPr></w:p>'
            new_para = etree.fromstring(new_para_str)
            sect_pr = new_para.find(f'.//{_W("sectPr")}')

            if ref_sect_pr is not None:
                for tag in ('w:pgSz', 'w:pgMar', 'w:cols'):
                    src = ref_sect_pr.find(_W(tag))
                    if src is not None:
                        sect_pr.append(deepcopy(src))

            body.insert(list(body).index(ch1_para), new_para)

    content['word/document.xml'] = etree.tostring(doc_root, xml_declaration=True, encoding='UTF-8', standalone=True)


def _set_pg_num_type(content, is_front_list):
    """按节设置 pgNumType"""
    doc_root = etree.fromstring(content['word/document.xml'])

    body = doc_root.find(_W('body'))
    all_sect = body.findall(f'.//{_W("sectPr")}')

    for i, sp in enumerate(all_sect):
        if i >= len(is_front_list):
            break
        # 删除旧的
        for old in sp.findall(_W('pgNumType')):
            sp.remove(old)
        # 创建新的
        pnt = etree.SubElement(sp, _W('pgNumType'))
        pnt.set(_W('start'), '1')
        if is_front_list[i]:
            pnt.set(_W('fmt'), 'upperRoman')

    content['word/document.xml'] = etree.tostring(doc_root, xml_declaration=True, encoding='UTF-8', standalone=True)


def _set_footer_distance(content, twips):
    """设置所有节的页脚距离（w:pgMar w:footer = twips）"""
    doc_root = etree.fromstring(content['word/document.xml'])
    body = doc_root.find(_W('body'))

    for sp in body.findall(f'.//{_W("sectPr")}'):
        pgmar = sp.find(_W('pgMar'))
        if pgmar is None:
            pgmar = etree.SubElement(sp, _W('pgMar'))
        pgmar.set(_W('footer'), str(twips))

    content['word/document.xml'] = etree.tostring(doc_root, xml_declaration=True, encoding='UTF-8', standalone=True)


def add_page_numbers(doc_path, output_path=None):
    doc_path = os.path.abspath(doc_path)
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f'文件不存在: {doc_path}')

    out_path = os.path.abspath(output_path) if output_path else doc_path

    print('处理文件:', doc_path)

    # ── 阶段一：读包 → XML 操作 → 写包 ──
    content = _read_zip(doc_path)

    # 1. 先清除所有旧页脚
    print('  清除旧页脚关系…')
    _clean_old_footers(content)

    # 2. 用 python-docx 找到正文第一章的段落索引
    ch1_idx, ch1_text = _find_ch1_start_index(content)
    if ch1_idx is not None:
        print(f'  正文第一章: 段落索引={ch1_idx}, text="{ch1_text[:30]}"')
    else:
        print('  未找到正文章节标题，全文按正文前处理')

    # 3. 修复分节符（清除旧 + 在 ch1_idx 前插入新）
    _fix_sections_in_package(content, ch1_idx)

    # 4. 写回并重新读取（计算节数）
    _write_zip(out_path, content)
    content2 = _read_zip(out_path)

    # 4. 用 python-docx 统计节数
    tmp_doc = Document(out_path)
    sec_count = len(tmp_doc.sections)
    del tmp_doc

    print(f'文档共 {sec_count} 节')

    # 5. 设置 pgNumType
    if ch1_idx is not None:
        # 第1节 = 正文前（upperRoman），末节 = 正文（decimal）
        is_front_list = [i < sec_count - 1 for i in range(sec_count)]
    else:
        is_front_list = [True] * sec_count

    _set_pg_num_type(content2, is_front_list)
    _set_footer_distance(content2, 851)  # 1.5cm = 851 twips

    for i in range(sec_count):
        label = '正文前' if is_front_list[i] else '正文'
        fmt = 'upperRoman' if is_front_list[i] else 'decimal'
        print(f'  第 {i+1} 节（{label}）→ pgNumType: {fmt}, start=1')

    # 6. 注入页脚
    for i in range(sec_count):
        is_front = is_front_list[i]
        label = '正文前' if is_front else '正文'

        if is_front and sec_count > 1:
            # 正文前：默认页脚（居中 PAGE）+ 首页页脚（空）+ 偶页页脚
            ftr_default = _build_footer_xml(centered=True, empty=False)
            _inject_footer(content2, i, 'default', ftr_default)
            ftr_first = _build_footer_xml(empty=True)
            _inject_footer(content2, i, 'first', ftr_first)
            _inject_footer(content2, i, 'even', ftr_default)
            print(f'    第 {i+1} 节（{label}）→ 默认页脚(PAGE), 首页页脚(空), 偶页页脚(PAGE)')
        else:
            # 正文：默认页脚（居中 PAGE）+ 偶页页脚
            ftr_default = _build_footer_xml(centered=True, empty=False)
            _inject_footer(content2, i, 'default', ftr_default)
            _inject_footer(content2, i, 'even', ftr_default)
            print(f'    第 {i+1} 节（{label}）→ 默认页脚(PAGE), 偶页页脚(PAGE)')

    # 写出最终结果
    _write_zip(out_path, content2)
    print(f'完成: {out_path}')


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    add_page_numbers(input_file, output_file)
