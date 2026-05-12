"""
patch_abstract.py - 修补中英文摘要的关键词格式

中文关键词："关键词：" 小4号黑体加粗，内容小4号宋体
英文关键词："Key words:" 小4号 Times New Roman 加粗，内容小4号 TNR 加粗
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys
import os

# ==================== 控制变量 ====================
FORMAT_KEYWORDS = True
# =================================================


def _make_run(paragraph, text, font_name='SimSun', font_size=Pt(12), bold=False, east_asia=None):
    """添加一个 run 到段落，设置字体字号"""
    run = paragraph.add_run(text)
    run.font.name = font_name
    ea = east_asia or font_name
    run._element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', ea
    )
    run.font.size = font_size
    run.font.bold = bold
    return run


def _clear_paragraph_runs(paragraph):
    """清除段落中的所有 run（保留 pPr）"""
    p = paragraph._element
    for child in list(p):
        if child.tag != qn('w:pPr'):
            p.remove(child)


def _set_first_line_indent(paragraph, indent_pt=Pt(24)):
    """设置首行缩进"""
    pf = paragraph.paragraph_format
    pf.first_line_indent = indent_pt


def _format_cn_keywords(para):
    """将中文关键词段落格式化为：'关键词：' 黑体加粗 + 内容宋体"""
    text = para.text.strip()
    if not text:
        return False

    _clear_paragraph_runs(para)
    _make_run(para, '关键词：', 'SimHei', Pt(12), bold=True, east_asia='SimHei')
    _make_run(para, text, 'SimSun', Pt(12), east_asia='SimSun')
    _set_first_line_indent(para)
    print(f'  已格式化中文关键词: "{text[:40]}..."')
    return True


def _format_en_keywords(para):
    """将英文关键词段落格式化为：'Key words: ' TNR 加粗 + 内容 TNR 加粗"""
    text = para.text.strip()
    if not text:
        return False

    _clear_paragraph_runs(para)
    _make_run(para, 'Key words: ', 'Times New Roman', Pt(12), bold=True, east_asia='SimSun')
    _make_run(para, text, 'Times New Roman', Pt(12), bold=True, east_asia='SimSun')
    _set_first_line_indent(para)
    print(f'  已格式化英文关键词: "{text[:40]}..."')
    return True


def patch_abstract(input_path, output_path):
    doc = Document(input_path)

    if not FORMAT_KEYWORDS:
        doc.save(output_path)
        return

    modified = False
    in_cn = False
    in_en = False
    cn_kw_para = None
    en_kw_para = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        cleaned = text.replace(' ', '').replace('　', '')

        # 进入中文摘要区域
        if cleaned == '摘要' and not in_cn and not in_en:
            in_cn = True
            cn_kw_para = None
            continue

        # 进入英文摘要区域
        if cleaned == 'Abstract':
            in_cn = False
            in_en = True
            en_kw_para = None
            continue

        # 遇到下一章标题 → 退出英文摘要
        if in_en and para.style.name.startswith('Heading'):
            in_en = False
            continue

        # 中文摘要内的段落：跟踪最后一个非空段落作为候选关键词
        if in_cn and text:
            cn_kw_para = para

        # 英文摘要内的段落：跟踪最后一个非空段落
        if in_en and text:
            en_kw_para = para

    # 格式化中文关键词（候选段落含中文分号才确认）
    if cn_kw_para is not None:
        text = cn_kw_para.text.strip()
        if '；' in text:
            modified |= _format_cn_keywords(cn_kw_para)

    # 格式化英文关键词
    if en_kw_para is not None:
        modified |= _format_en_keywords(en_kw_para)

    if modified:
        doc.save(output_path)
        print(f'处理完成，已保存至: {output_path}')
    else:
        print('未找到需要格式化的关键词段落')
        doc.save(output_path)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('用法: python patch_abstract.py <输入文件> [输出文件]')
        sys.exit(1)

    input_file = sys.argv[1]
    if len(sys.argv) > 2:
        output_file = sys.argv[2]
    else:
        base, ext = os.path.splitext(input_file)
        output_file = f'{base}_abstract{ext}'

    if not os.path.exists(input_file):
        print(f'错误: 输入文件不存在 - {input_file}')
        sys.exit(1)

    patch_abstract(input_file, output_file)
