"""
patch_header.py - 使用 python-docx 添加页眉（仿 99.docx 模板）

页眉内容：
  - 居中文字：小五号宋体
  - 页眉底线（段落底部边框）
  - 右侧浮动 LOGO（绝对定位，相对页边距）

用法：python patch_header.py <输入文件> [输出文件]

依赖：python-docx (pip install python-docx)
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import sys


# ==================== 控制变量 ====================
HEADER_TEXT = "latex2word-thesis 构建产物,修改LOGO.png和patch_header修改信息"  # 页眉文字

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(_SCRIPT_DIR, "figures", "LOGO.png")
if not os.path.exists(LOGO_PATH):
    LOGO_PATH = os.path.join(_SCRIPT_DIR, "LOGO.png")
# =================================================

XIAO_WU = 9          # 小五号 = 9pt
IMG_SIZE = Cm(2)     # 2cm

# 浮动图片位置（相对页边距）
LOGO_LEFT_EMU = int(16.55 / 2.54 * 914400)  # 16.55cm → EMU
LOGO_TOP_EMU = int(-1.71 / 2.54 * 914400)   # -1.71cm → EMU
# 四周型环绕距正文距离
IMG_WRAP_DIST = str(int(0.32 / 2.54 * 914400))  # 0.32cm → EMU


def _add_bottom_border(paragraph):
    """给段落加底部边框（single, 0.75pt, auto color）"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    for attr, val in [('val', 'single'), ('sz', '6'), ('space', '0'), ('color', 'auto')]:
        bottom.set(qn('w:' + attr), val)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_floating_picture(paragraph, image_path):
    """添加浮动图片（绝对定位，相对页边距）"""
    run = paragraph.add_run()
    pic = run.add_picture(image_path, width=IMG_SIZE, height=IMG_SIZE)

    drawing = run._r.find(qn('w:drawing'))
    if drawing is None:
        return

    inline = drawing.find(qn('wp:inline'))
    if inline is None:
        return

    # 创建 anchor 元素替换 inline
    anchor = OxmlElement('wp:anchor')
    for attr, val in [('distT', '0'), ('distB', '0'),
                      ('distL', IMG_WRAP_DIST), ('distR', IMG_WRAP_DIST),
                      ('simplePos', '0'), ('relativeHeight', '0'),
                      ('behindDoc', '0'), ('locked', '0'),
                      ('layoutInCell', '1'), ('allowOverlap', '1')]:
        anchor.set(attr, val)

    # simplePos
    sp = OxmlElement('wp:simplePos')
    sp.set('x', '0')
    sp.set('y', '0')
    anchor.append(sp)

    # positionH (水平，相对页边距)
    posH = OxmlElement('wp:positionH')
    posH.set('relativeFrom', 'margin')
    offH = OxmlElement('wp:posOffset')
    offH.text = str(LOGO_LEFT_EMU)
    posH.append(offH)
    anchor.append(posH)

    # positionV (垂直，相对页边距)
    posV = OxmlElement('wp:positionV')
    posV.set('relativeFrom', 'margin')
    offV = OxmlElement('wp:posOffset')
    offV.text = str(LOGO_TOP_EMU)
    posV.append(offV)
    anchor.append(posV)

    # extent（图片尺寸）
    ext = OxmlElement('wp:extent')
    ext.set('cx', inline.find(qn('wp:extent')).get('cx'))
    ext.set('cy', inline.find(qn('wp:extent')).get('cy'))
    anchor.append(ext)

    # effectExtent
    eff = inline.find(qn('wp:effectExtent'))
    if eff is not None:
        anchor.append(eff)

    # wrapNone（不和文字混排）
    wrap = OxmlElement('wp:wrapNone')
    anchor.append(wrap)

    # 复制剩余子元素（docPr, cNvGraphicFramePr, graphic 等）
    for child in list(inline):
        tag = child.tag
        if tag in [qn('wp:extent'), qn('wp:effectExtent')]:
            continue
        anchor.append(child)

    # 替换 inline 为 anchor
    drawing.replace(inline, anchor)


def _setup_header(header):
    """清空页眉保留一个空段落，链接到前一节"""
    try:
        header.is_linked_to_previous = True
    except Exception:
        pass

    # 删除多余段落，只保留第一个
    while len(header.paragraphs) > 1:
        p = header.paragraphs[-1]
        p._p.getparent().remove(p._p)

    # 清空第一个段落
    p = header.paragraphs[0]
    p.clear()
    return p


def _write_header(header):
    """写页眉文字 + 底线 + LOGO"""
    p = _setup_header(header)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(HEADER_TEXT)
    run.font.name = '宋体'
    run.font.size = Pt(XIAO_WU)
    # 设置东亚字体（宋体），否则 Word 默认用等线
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    _add_bottom_border(p)

    if os.path.exists(LOGO_PATH):
        _add_floating_picture(p, LOGO_PATH)
    else:
        print(f"    警告: LOGO 不存在 - {LOGO_PATH}")


def _has_title_pg(section):
    """检测节是否有首页页眉（w:titlePg）"""
    sp = section._sectPr
    return sp.find(qn('w:titlePg')) is not None


def _has_even_header(section):
    """检测节是否有偶页页眉（w:evenAndOddHeaders）"""
    sp = section._sectPr
    return sp.find(qn('w:evenAndOddHeaders')) is not None


def add_header(doc_path, output_path=None):
    doc_path = os.path.abspath(doc_path)
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"文件不存在: {doc_path}")

    out_path = os.path.abspath(output_path) if output_path else doc_path

    print("处理文件:", doc_path)
    print(f"页眉文字: {HEADER_TEXT}")
    print(f"LOGO 路径: {LOGO_PATH}")

    doc = Document(doc_path)

    for i, section in enumerate(doc.sections):
        print(f"  处理第 {i + 1} 节...")

        # 只处理文档已有的页眉类型，不改 section 属性
        # （不设 different_first_page_header_footer，避免打乱页脚结构）
        header_types = [('header', '默认')]

        if _has_title_pg(section):
            header_types.append(('first_page_header', '首页'))
        if _has_even_header(section):
            header_types.append(('even_page_header', '偶页'))

        for attr_name, label in header_types:
            try:
                header = getattr(section, attr_name)
            except Exception:
                continue
            _write_header(header)
            print(f"    第 {i + 1} 节 ({label}) 页眉已添加")

    doc.save(out_path)
    print(f"完成: {out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.exists(LOGO_PATH):
        print(f"警告: LOGO 文件未找到 ({LOGO_PATH})，将继续但不插入图片")

    add_header(input_file, output_file)
